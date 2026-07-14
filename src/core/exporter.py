import io
from fpdf import FPDF
from docx import Document

def generate_pdf(prompt_text: str) -> bytes:
    """
    Generates a PDF from the given prompt text.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    # Title
    pdf.set_font("Arial", style='B', size=16)
    pdf.cell(200, 10, txt="Optimized Prompt", ln=True, align='C')
    pdf.ln(10)
    
    # Body
    pdf.set_font("Arial", size=12)
    # multi_cell handles line wrapping
    pdf.multi_cell(0, 10, txt=prompt_text)
    
    # Return as bytes
    return pdf.output(dest='S').encode('latin1', 'ignore')

def generate_docx(prompt_text: str) -> bytes:
    """
    Generates a DOCX from the given prompt text.
    """
    doc = Document()
    doc.add_heading('Optimized Prompt', 0)
    doc.add_paragraph(prompt_text)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
